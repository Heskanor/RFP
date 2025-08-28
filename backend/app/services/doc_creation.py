from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from bs4 import BeautifulSoup
import datetime
from app.models.routers_models import ExportParams
from app.config.firebase import firebase_manager
from app.models.models import Collections
import tempfile
import os
from uuid import uuid4
from app.models.models import File
from collections import defaultdict
import traceback
class DocCreationService:
    def __init__(self):
        self.doc = Document()

    def add_heading(self, text, level=1):
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text, bold=False, italic=False, size=12):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return para

    def add_bullet_point(self, text, level=0):
        """
        Add a bullet point with optional nesting level.
        
        :param doc: The Document instance
        :param text: Text of the bullet point
        :param level: Indentation level (0 = top level)
        """
        para = self.doc.add_paragraph(text, style='List Bullet')
        if level > 0:
            para.paragraph_format.left_indent = Pt(14 * level)
        return para

    def add_table(self, data, col_names):
        table = self.doc.add_table(rows=1, cols=len(col_names))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(col_names):
            hdr_cells[i].text = name

        for row in data:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
                
    def add_hyperlink(self, paragraph, url, label, color="0000FF", underline=True):
        """
        Add a clickable hyperlink to a paragraph.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: The target URL.
        :param label: The text to display.
        :param color: The font color (default: blue).
        :param underline: Whether the link should be underlined.
        """
        # Create the w:hyperlink tag and add needed attributes
        part = paragraph.part
        r_id = part.relate_to(url,RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create the w:r element
        new_run = OxmlElement("w:r")

        # Create a w:rPr element
        rPr = OxmlElement("w:rPr")

        # Style: color
        color_elem = OxmlElement("w:color")
        color_elem.set(qn("w:val"), color)
        rPr.append(color_elem)

        # Style: underline
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)

        new_run.append(rPr)

        # Text element
        text = OxmlElement("w:t")
        text.text = label
        new_run.append(text)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return paragraph
    
    def add_html_answer(self, html):
        soup = BeautifulSoup(html, 'html.parser')
        references = []  # Store ref-tags with their refid

        for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'strong', 'span']):
            if elem.name == 'p':
                para = self.doc.add_paragraph()
                for child in elem.children:
                    if child.name == 'span' and 'ref-tag' in child.get('class', []):
                        refid = child.get('data-ref-id') or child.get_text()
                        run = para.add_run(f"[ref: {refid.split('_')[-1]}]")  # You can style this if needed
                        run.font.highlight_color = 3  # Highlight yellow
                        references.append(refid)
                    elif hasattr(child, 'name'):
                        run = para.add_run(child.get_text())
                    else:
                        run = para.add_run(str(child))
                    run.font.size = Pt(11)

            elif elem.name in ['h1', 'h2', 'h3']:
                heading_level = {"h1": 4, "h2": 5, "h3": 6}.get(elem.name, 4)
                self.doc.add_heading(elem.get_text(), level=heading_level)

            elif elem.name == 'strong':
                run = self.doc.add_paragraph().add_run(elem.get_text())
                run.bold = True

        return references

    def save_docx_temp(self) -> str:
        temp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        self.doc.save(temp.name)
        return temp.name

    async def create_doc(self, project_id: str, export_params: ExportParams):
        try:
            self.doc = Document()  # Ensure a fresh document on each call

            project = await firebase_manager.get_document(Collections.PROJECT, project_id)
            if not project:
                raise ValueError("Project not found")

            user_id = project.get("user_id", "")
            details = project.get("details", {})

            # Title Page (No page break here)
            self.doc.add_heading(f"Proposal Draft: {project.get('title', '')}", 0)
            self.add_paragraph(f"Generated on {datetime.date.today().isoformat()}", italic=True)

            # Summary Section
            if export_params.summary:
                self.doc.add_page_break()
                self.add_heading("Executive Summary", level=1)
                self.add_paragraph(details.get("summary", ""))

            # Stakeholders
            if export_params.stakeholders:
                self.add_heading("Stakeholders", level=1)
                if details.get("stakeholders", []):
                    self.add_table([(s['name'], s['title'], s['email']) for s in details.get("stakeholders", [])], ['Name', 'Title', 'Email'])

            # Timeline
            if export_params.timeline and details.get("timeline", []):
                self.add_heading("Timeline", level=1)
                self.add_table([(t['date'], t['milestone']) for t in details.get("timeline", [])], ['Date', 'Milestone'])

            # Tickets
            if export_params.tickets:
                if export_params.filteredTickets:
                    tickets = await firebase_manager.get_documents(Collections.TICKET, list_ids=export_params.filteredTickets)
                else:
                    tickets = await firebase_manager.query_collection(Collections.TICKET.value, filters=[("project_id", "==", project_id)])

                if tickets:
                    # Group tickets by their "type"
                    grouped_tickets = defaultdict(list)
                    for ticket in tickets:
                        ticket_type = ticket.get('type', 'Other').capitalize()
                        grouped_tickets[ticket_type].append(ticket)

                    for ticket_type, typed_tickets in grouped_tickets.items():
                        self.doc.add_page_break()
                        self.add_heading(ticket_type, level=1)

                        for ticket in typed_tickets:
                            self.add_heading(ticket.get('title', ''), level=2)
                            self.add_paragraph(ticket.get('description', ''))
                            self.add_heading("Answer", level=3)

                            if isinstance(ticket.get('answer', {}), dict):
                                refs = self.add_html_answer(ticket.get('answer', {}).get("markdown", ''))
                                if refs:
                                    self.add_heading("References", level=4)
                                    for ref in refs:
                                        ref_sources = ticket.get('answer', {}).get("references", {}).get(ref, {}).get("sources", [])
                                        if ref_sources:
                                            for source in ref_sources:
                                                para = self.add_bullet_point(f"Reference - {ref.split('_')[-1]}: [Pages: {tuple(source.get('pages'))}] - ", level=1)
                                                self.add_hyperlink(para, source.get("url"), source.get("text"))


            # path = self.doc.save(f"temp/Proposal Draft - {project.get('title', '')}.docx")
            # return path
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                self.doc.save(tmp.name)
                tmp_path = tmp.name
            file_id = str(uuid4())
            public_url = await firebase_manager.upload_file(
                file=None,  
                path_segments=[user_id, project_id],
                file_id=file_id,
                file_source=tmp_path,  # this triggers upload_from_filename
                file_type="docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            file = File(
                id=file_id,
                user_id=user_id,
                project_id=project_id,
                is_knowledge_hub=False,
                is_supporting_file=False,
                is_proposal_draft=True,
                name=f"Proposal Draft - {project.get('title', '')}",
                type="docx",
                url=public_url,         
            )
            await firebase_manager.create_document(Collections.FILE, file, file_id)
            os.remove(tmp_path)
            return {"success": True, "file_id": file_id}
        except Exception as e:
            print(f"Error creating doc: {e}")
            traceback.print_exc()
            return {"success": False, "error": str(e)}

doc_creation_service = DocCreationService()